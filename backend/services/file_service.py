import os
from typing import List
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings.fastembed import FastEmbedEmbeddings
from langchain_chroma import Chroma
from PIL import Image
import io
import pdfplumber
import docx
from backend.core.config import settings
import logging

# Setup logger
logger = logging.getLogger("uvicorn.error")

MAX_TEXT_CHARS = 30000 # Increased limit

class FileService:
    def __init__(self):
        logger.info("Initializing FileService with FastEmbed (CPU-optimized)")
        self.embeddings = FastEmbedEmbeddings(model_name="BAAI/bge-small-en-v1.5")
        self.vector_store = Chroma(
            persist_directory=settings.CHROMA_PERSIST_DIR,
            embedding_function=self.embeddings
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.EMBED_CHUNK_SIZE,
            chunk_overlap=50
        )
        self.upload_dir = "backend/uploads"
        os.makedirs(self.upload_dir, exist_ok=True)

    async def ingest_file(self, file_content: bytes, filename: str):
        """Ingest a file (PDF, Text, Code, Image, Docx) into the vector store."""
        logger.info(f"Ingesting file: {filename}")
        # Save file to disk for static serving
        file_path = os.path.join(self.upload_dir, filename)
        with open(file_path, "wb") as f:
            f.write(file_content)
            
        ext = os.path.splitext(filename)[1].lower()
        
        try:
            from langchain_core.documents import Document
            text = ""
            
            if ext == ".pdf":
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                logger.info(f"Extracted {len(text)} characters from PDF {filename}")
                    
            elif ext == ".docx":
                doc = docx.Document(io.BytesIO(file_content))
                text = "\n".join([para.text for para in doc.paragraphs])
                logger.info(f"Extracted {len(text)} characters from DOCX {filename}")
                
            elif ext in [".txt", ".md", ".py", ".js", ".ts", ".tsx", ".html", ".css", ".json", ".lock"]:
                try:
                    text = file_content.decode("utf-8")
                except UnicodeDecodeError:
                    text = file_content.decode("utf-8", errors="ignore")
                logger.info(f"Decoded {len(text)} characters from text file {filename}")
                
            elif ext in [".png", ".jpg", ".jpeg", ".webp"]:
                # For images, we don't extract text for now (slowness fix)
                logger.info(f"Image file {filename} saved. Skipping text extraction.")
                return file_path
            else:
                # Default to text decoding if it looks like text
                try:
                    text = file_content.decode("utf-8")
                    logger.info(f"Decoded {len(text)} characters from unknown file type {filename}")
                except UnicodeDecodeError:
                    logger.warning(f"File {filename} appears to be binary. Skipping ingestion.")
                    return file_path
                
            if not text.strip() or len(text.strip()) < 5:
                logger.warning(f"Extracted text from {filename} is too short or empty. Skipping vector store.")
                return file_path

            # Truncate text if it's too long
            if len(text) > MAX_TEXT_CHARS:
                text = text[:MAX_TEXT_CHARS] + "\n\n[Note: Document truncated]"

            documents = [Document(page_content=text, metadata={"source": filename})]
            texts = self.text_splitter.split_documents(documents)
            
            if texts:
                logger.info(f"Adding {len(texts)} chunks to vector store for {filename}")
                # Log first chunk to verify
                logger.info(f"Sample chunk: {texts[0].page_content[:100]}...")
                self.vector_store.add_documents(texts)
                logger.info(f"Successfully added {filename} to vector store.")
            else:
                logger.warning(f"No text chunks generated for {filename}.")
            
            return file_path
        except Exception as e:
            logger.error(f"Error ingesting file {filename}: {e}")
            import traceback
            traceback.print_exc()
            return None

    def get_retriever(self):
        # Always return a fresh retriever from the current vector store
        return self.vector_store.as_retriever(search_kwargs={"k": 10})

    def reset_vector_store(self):
        """Clears the vector store safely by deleting all documents."""
        logger.info("Resetting vector store...")
        try:
            # Delete all documents in the collection
            all_docs = self.vector_store.get()
            if all_docs and all_docs['ids']:
                self.vector_store.delete(ids=all_docs['ids'])
                logger.info(f"Deleted {len(all_docs['ids'])} documents from vector store.")
            
            # Also clear uploads directory
            if os.path.exists(self.upload_dir):
                for f in os.listdir(self.upload_dir):
                    file_path = os.path.join(self.upload_dir, f)
                    try:
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                    except Exception as e:
                        logger.warning(f"Could not remove file {file_path}: {e}")
            
            logger.info("Vector store reset successfully.")
            return True
        except Exception as e:
            logger.error(f"Error resetting vector store: {e}")
            # Fallback: try to re-initialize
            try:
                self.vector_store = Chroma(
                    persist_directory=settings.CHROMA_PERSIST_DIR,
                    embedding_function=self.embeddings
                )
                return True
            except:
                return False

file_service = FileService()
